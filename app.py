import streamlit as st
from openai import OpenAI
from dotenv import load_dotenv
import os
import PyPDF2
import docx
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# ---------------- LOAD API ----------------
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

st.title("📄 AI Resume Improver 🚀")

st.write("Upload your resume (PDF, DOCX, TXT) or paste text.")

# ---------------- FILE UPLOAD ----------------
uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx", "txt"])

resume_text = ""

if uploaded_file is not None:
    file_type = uploaded_file.name.split(".")[-1]

    if file_type == "pdf":
        reader = PyPDF2.PdfReader(uploaded_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                resume_text += text

    elif file_type == "docx":
        doc = docx.Document(uploaded_file)
        for para in doc.paragraphs:
            resume_text += para.text + "\n"

    elif file_type == "txt":
        resume_text = str(uploaded_file.read(), "utf-8")

# ---------------- TEXT INPUT ----------------
resume_text = st.text_area("Resume Text:", value=resume_text, height=300)

# ---------------- SESSION STATE ----------------
if "resume_output" not in st.session_state:
    st.session_state.resume_output = ""

# ---------------- IMPROVE RESUME ----------------
if st.button("✨ Improve Resume"):
    if resume_text.strip() != "":
        with st.spinner("Improving your resume..."):

            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {
                        "role": "system",
                        "content": """You are an expert resume reviewer.

Improve the resume professionally.

Make it:
- ATS-friendly
- Clear and impactful
- Use strong action words

Also provide:
1. Improved resume version
2. 5 improvement suggestions
"""
                    },
                    {
                        "role": "user",
                        "content": resume_text
                    }
                ]
            )

            result = response.choices[0].message.content
            st.session_state.resume_output = result

    else:
        st.warning("Please upload or paste a resume.")

# ---------------- SHOW RESULT ----------------
if st.session_state.resume_output:
    st.subheader("✅ Improved Resume")
    st.write(st.session_state.resume_output)

    # ---------------- CUSTOM CHAT ----------------
    st.subheader("💬 Customize Your Resume")

    user_msg = st.text_input("Modify your resume (e.g., 'Make it more professional')")

    if user_msg:
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a resume assistant. Modify resume based on user request."
                },
                {
                    "role": "user",
                    "content": f"Resume:\n{st.session_state.resume_output}\n\nUser Request:\n{user_msg}"
                }
            ]
        )

        updated = response.choices[0].message.content
        st.session_state.resume_output = updated

        st.write("🔄 Updated Resume:")
        st.write(updated)

    # ---------------- DOWNLOAD WORD ----------------
    if st.button("📥 Download as Word"):
        doc = Document()
        doc.add_paragraph(st.session_state.resume_output)
        doc.save("resume.docx")

        with open("resume.docx", "rb") as file:
            st.download_button(
                label="Download Word File",
                data=file,
                file_name="resume.docx"
            )

    